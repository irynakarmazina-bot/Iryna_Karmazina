#!/usr/bin/env python3
# Генерація документів за бланками з автоприкріпленням до угоди.
# Типи: loi (LOI CMA), zayavka_cma (Заявка CMA), zayavka (Заявка на перевезення),
#       kkk (Заявка Кампус Коттон Клаб), t1 (Заявка на Т1).
import http.server, socketserver, json, urllib.request, urllib.error, urllib.parse, os, io, re, uuid, datetime
import docx
PORT = 8790
NC = "http://localhost:8080"
USERS_T = "meqpi0r197bz14n"; DISP_T = "m58xsjo6at01ohl"
TPL_DIR = os.environ.get("TPL_DIR", "/root/doc-templates")
DENY_ROLES = {"Перегляд", "Логіст", None}
_TOK = None
def admin_tok():
    global _TOK
    if _TOK is None:
        _TOK = open("/root/nocodb-token.txt").read().strip()
    return _TOK

MONTHS = ["січня","лютого","березня","квітня","травня","червня","липня","серпня","вересня","жовтня","листопада","грудня"]
def date_ukr(iso):
    try:
        d = datetime.date.fromisoformat(iso)
        return "«%d» %s %d року" % (d.day, MONTHS[d.month-1], d.year)
    except Exception:
        return iso
def date_dot(iso):
    try:
        return datetime.date.fromisoformat(iso).strftime("%d.%m.%Y")
    except Exception:
        return iso

def sval(fields, key): return str(fields.get(key, "") or "")
def fmt(tpl, fields):
    class D(dict):
        def __missing__(self, k): return ""
    return tpl.format_map(D({k: str(v) for k, v in fields.items()}))

def all_paras(d):
    out = list(d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for c in row.cells:
                out += c.paragraphs
    return out

def set_para(p, text):
    lines = str(text).split("\n")
    if p.runs:
        p.runs[0].text = lines[0]
        for r in p.runs[1:]: r.text = ""
    else:
        p.add_run(lines[0])
    for ln in lines[1:]:
        p.runs[0].add_break()
        p.add_run(ln)

def set_cell(cell, text):
    set_para(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        for r in p.runs: r.text = ""

# ── LOI CMA та Заявка CMA: заміна зразкових значень ──
LOI_MAP = [
  ("«EXPORT INTERIOR» LLC, Enterprise code 46290240, 03035, Kyiv, Solomyanska Square, building No. 2, office 907, export.interior.is@gmail.com, tel. +38 (096) 009-55-01", "shipper"),
  ("Company «F. Engey EHF», Sundabakka 4, 340 Stykkisholmur,Iceland, Reg. number:  660489-1189, Tax VAT: 17307, engey@simnet.is, tel: +35 489 38 171", "consignee"),
  ("Place of Receipt: Kyiv", "!Place of Receipt: {place_receipt}"),
  ("Reykjavik", "place_delivery"),
  ("0C47VW1MA", "voyage"),
  ("OEA0218983", "booking"),
  ("1x40HC", "container"),
  ("FURNITURE, DOORS", "goods"),
  ("3.06.26", "date"),
]
ZAYCMA_MAP = [
  ("ЗАЯВКА № 1", "!ЗАЯВКА № {req_no}"),
  ("Адреса / Address: Україна, 03035, місто Київ, Солом’янська площа, будинок 2, офіс 907", "!Адреса / Address: {shipper_addr}"),
  ("Адреса / Address: Sundabakka 4, 340 Stykkisholmur,Iceland", "!Адреса / Address: {consignee_addr}"),
  ("ТОВ «ЕКСПОРТ ІНТЕРІОР»/ «EXPORT INTERIOR» LLC", "shipper_name"),
  ("Код ЄДРПОУ 46290240", "!Код ЄДРПОУ {shipper_edrpou}"),
  ("Company «F. Engey EHF»", "consignee_name"),
  ("Олега Мудрака, 1, літера Ш, https://maps.app.goo.gl/3WRMcvHUj77a4tHx8?g_st=iv", "loading_addr"),
  ("+380960095501 Юлія", "loading_contact"),
  ("Малинська вул, 20, Київ, Наталія Сергіївна +380636912691, +380675053108", "customs_addr"),
  ("5.06.26, о 9 ранку", "loading_dt"),
  ("Митниця працює до 17.00", "customs_note"),
  ("Навантажувач, ручне завантаження", "load_method"),
  ("Київ - Гданськ", "route"),
  ("OEA0218983", "booking"),
  ("1х40НС", "container"),
  ("17 т", "weight"),
  ("9403401000 кухні", "goods1"),
  ("9403601000 інші меблі", "goods2"),
  ("4418298000 двірне полотно", "goods3"),
  ("Палети та коробки", "packing"),
  ("1750 USD", "rate"),
  ("3.06.2026", "date"),
]

def fill_map(tpl, mapping, fields):
    """Заповнює бланк, замінюючи ЗРАЗКОВИЙ текст реальної минулої угоди.

    НЕБЕЗПЕКА ЦЬОГО ПІДХОДУ (знайдено при аудиті 02.08.2026): у шаблонах зашиті
    справжні реквізити попереднього клієнта — назва, код ЄДРПОУ, пошта, букінг.
    Якщо зразок у шаблоні не знайдено (хтось відкрив .docx у Word і поправив
    літеру, або Word сам розбив рядок), заміна просто не відбувалась — і в
    готовий документ ішли ЧУЖІ реквізити, а сервер рапортував «успіх».
    Такий документ пішов би в лінію від імені Юнітекса.

    Тепер: перед заповненням звіряємо, що КОЖЕН зразок є в шаблоні. Немає хоч
    одного — документ не створюється взагалі, з поясненням, що саме не знайдено.
    Краще не видати документ, ніж видати з чужими даними.
    """
    d = docx.Document(os.path.join(TPL_DIR, tpl))
    paras = list(all_paras(d))
    # знімок ДО замін: інакше зразок, з'їдений попередньою заміною,
    # виглядав би як «відсутній у шаблоні»
    whole = "\n".join(p.text for p in paras)
    missing = [key.lstrip("!") for sample, key in mapping if sample not in whole]
    if missing:
        raise RuntimeError(
            "Шаблон %s змінено: не знайдено зразки для полів %s%s. Документ НЕ створено, "
            "щоб не підставити реквізити іншого клієнта." %
            (tpl, ", ".join(missing[:3]), " та ін." if len(missing) > 3 else ""))
    for p in paras:
        t = p.text; orig = t
        for sample, key in mapping:
            if sample in t:
                val = fmt(key[1:], fields) if key.startswith("!") else sval(fields, key)
                t = t.replace(sample, val)
        if t != orig:
            set_para(p, t)
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()

# ── Заявка на перевезення (додаток до договору ТЕО) ──
ZP_ROWS = {  # рядок таблиці → ключ поля (колонка 3)
  0: "services", 1: "route", 2: "container", 3: "con_bl", 4: "uktzed",
  5: "shipper", 6: "consignee", 7: "incoterms", 8: "eta",
  9: "pol", 10: "pod", 11: "customs_addr", 12: "unload_addr", 13: "price", 18: "other",
}
ZP_TEXT = [
  ("№ 12-30.09.24 від 30.09.2024 року", "contract"),
  ("ЗАЯВКА № 7", "!ЗАЯВКА № {num}"),
  ("«5» травня 2026 року", "date_ukr"),
  ("ТОВАРИСТВО З ОБМЕЖЕНОЮ ВІДПОВІДАЛЬНІСТЮ ТД «РТК»", "client_full"),
  ("Новікова Олексія Олексійовича", "client_dir_gen"),
  ("ТОВ ТД «РТК»", "client_short"),
  ("Код ЄДРПОУ 41572243", "!Код ЄДРПОУ {client_edrpou}"),
  ("О.О. Новіков", "client_dir_short"),
]
def fill_zp(fields):
    if fields.get("date") and not fields.get("date_ukr"):
        fields["date_ukr"] = date_ukr(fields["date"])
    d = docx.Document(os.path.join(TPL_DIR, "zayavka_perevezennya.docx"))
    for p in all_paras(d):
        t = p.text; orig = t
        for sample, key in ZP_TEXT:
            if sample in t:
                val = fmt(key[1:], fields) if key.startswith("!") else sval(fields, key)
                t = t.replace(sample, val)
        if t != orig:
            set_para(p, t)
    tbl = d.tables[0]
    for ri, key in ZP_ROWS.items():
        v = fields.get(key)
        if v is not None and str(v) != "":
            set_cell(tbl.rows[ri].cells[2], v)
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()

# ── Заявка Кампус Коттон Клаб: перезапис абзаців за міткою ──
KKK_LABELS = [
  ("Пункт відправлення:", "pickup"),
  ("Готовність вантажу до відправлення:", "ready_date"),
  ("Пункт прибуття:", "destination"),
  ("Тип транспорту та маршрут міжнародного перевезення:", "transport_route"),
  ("Митниця в країні відправлення", "customs_from"),
  ("Пункт перетину кордону", "border"),
  ("Митниця в країні прибуття", "customs_to"),
  ("Кінцевий пункт доставки:", "final_addr"),
  ("Тип авто:", "truck_type"),
  ("Вага вантажу (брутто):", "weight_cargo"),
  ("Загальна вартість міжнародного перевезення:", "price_text"),
]
KKK_TEXT = [
  ("№10-02-2025", "!№{num}"),
  ("«10» лютого 2025 року", "date_ukr"),
  ("(номер замовлення: б/н)", "!(номер замовлення: {order_no})"),
]
def fill_kkk(fields):
    if fields.get("date") and not fields.get("date_ukr"):
        fields["date_ukr"] = date_ukr(fields["date"])
    if not fields.get("order_no"): fields["order_no"] = "б/н"
    d = docx.Document(os.path.join(TPL_DIR, "kkk_zayavka.docx"))
    for p in all_paras(d):
        t = p.text
        done = False
        for label, key in KKK_LABELS:
            v = fields.get(key)
            if t.strip().startswith(label) and v is not None and str(v) != "":
                sep = "" if label.endswith(":") else " -"
                set_para(p, "%s%s %s" % (label, sep, v))
                done = True
                break
        if done: continue
        orig = t
        for sample, key in KKK_TEXT:
            if sample in t:
                val = fmt(key[1:], fields) if key.startswith("!") else sval(fields, key)
                t = t.replace(sample, val)
        if t != orig:
            set_para(p, t)
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()

# ── Заявка на Т1: заповнення пропусків (___) після мітки ──
T1_LABELS = [
  ("Держ. номер тягача", "truck"),
  ("Держ. номер причепа", "trailer"),
  ("Запланований прикордонний перехід", "border"),
  ("Орієнтовні дата та час прибуття", "eta"),
  ("Номер телефону водія", "driver_phone"),
  ("Відповідальний за розмитнення", "broker"),
  ("Відповідальний за подачу заявки", "submitter"),
  ("Вартість вантажу за інвойсом", "invoice_value"),
  ("Кількість кодів УКТЗЕД", "codes_count"),
  ("Місце відкриття Т1", "open_place"),
  ("Характер вантажу", "cargo_nature"),
  ("Маршрут (звідки – куди):", "route"),
  ("Дата заявки", "date_txt"),
]
def fill_t1(fields):
    if fields.get("date") and not fields.get("date_txt"):
        fields["date_txt"] = date_dot(fields["date"])
    d = docx.Document(os.path.join(TPL_DIR, "t1_zayavka.docx"))
    for p in all_paras(d):
        t = p.text
        if "___" not in t: continue
        for label, key in T1_LABELS:
            v = fields.get(key)
            if label in t and v is not None and str(v) != "":
                set_para(p, re.sub(r"_{3,}", "  " + str(v) + "  ", t, count=1))
                break
    # клітинка "Дата заявки" в таблиці (мітка й пропуск у різних клітинках)
    for tb in d.tables:
        for row in tb.rows:
            cells = row.cells
            for i, c in enumerate(cells[:-1]):
                if c.text.strip().startswith("Дата заявки") and "___" in cells[i+1].text and fields.get("date_txt"):
                    set_cell(cells[i+1], fields["date_txt"])
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()

# ── Maersk POA / LOI (вивіз: ми або інший експедитор — компанія задається полем) ──
MPOA_MAP = [
  ("15.07.2026", "date_txt"),
  ("MRKU2438561", "container"),
  ("270859899for", "!{bl} for"),
  ("270859899", "bl"),
  ("LLC UNITEX HD (65076, ODESSA REGION, ODESSA, UKRAINE, OFFICE 56 RADISNA STR. 3)", "agent"),
]
def fill_mpoa(fields):
    if fields.get("date") and not fields.get("date_txt"):
        fields["date_txt"] = date_dot(fields["date"])
    return fill_map("maersk_poa.docx", MPOA_MAP, fields)

MLOI_MAP = [
  ("15.07.2026", "date_txt"),
  ("270859899", "booking"),
  ("WIKING / 627N", "vessel_voyage"),
  ("CELLULOSE ETHER", "goods"),
  ("MRKU2438561", "containers"),
]
def fill_mloi(fields):
    if fields.get("date") and not fields.get("date_txt"):
        fields["date_txt"] = date_dot(fields["date"])
    return fill_map("maersk_loi.docx", MLOI_MAP, fields)

# ── Заявка на авто в Maersk (бланк .xlsx: колонка A — мітка, колонка B — значення) ──
# Тут, на відміну від docx-бланків, шукаємо не зразкове ЗНАЧЕННЯ, а МІТКУ в колонці A,
# і завжди перезаписуємо колонку B — навіть коли поле порожнє. Інакше в готовій заявці
# лишились би дані клієнта, з якого знято бланк (адреса заводу, телефон водія, ставка).
MZAY_ROWS = [
  ("ВАЖНО", "note"),
  ("Дата и время подачи", "pickup_dt"),
  ("НОМЕР БУКИНГА", "booking"),
  ("КОЛ-ВО КОНТЕЙНЕРОВ", "cnt"),
  ("ТИП КОНТЕЙНЕРА", "cont_type"),
  ("СУДОХОДНАЯ ЛИНИЯ", "line"),
  ("ГРУЗ", "goods"),
  ("ВЕС БРУТТО ГРУЗА", "weight"),
  ("МЕСТО ПОГРУЗКИ,КОНТАКТ", "loading_addr"),
  ("ТАМОЖНЯ,КОНТАКТ", "customs_addr"),
  ("Особые требования", "special"),
  ("Порт погрузки", "pol"),
  ("Порт перевалки", "transship"),
  ("Порт выгрузки", "pod"),
  ("Судно", "vessel_voyage"),
  ("Дата судозахода", "eta"),
  ("Оговоренная ставка", "rate"),
  ("Валюта платежа", "currency"),
  ("Код компании в системе Маерска", "maersk_code"),
  ("НАЗВАНИЕ Компании в системе Маерска", "maersk_name"),
]
# мітки немає, значення лежить у рядку нижче тієї ж колонки (контакт на завантаженні)
MZAY_NEXT = {"loading_addr": "loading_contact"}

def _norm(s): return re.sub(r"\s+", " ", str(s if s is not None else "")).strip().lower()

def _row_by_label(ws, label):
    n = _norm(label)
    for r in range(1, ws.max_row + 1):
        if _norm(ws.cell(row=r, column=1).value).startswith(n):
            return r
    return None

def fill_mzay(fields):
    import openpyxl
    wb = openpyxl.load_workbook(os.path.join(TPL_DIR, "maersk_zayavka.xlsx"))
    ws = wb.worksheets[0]
    rows = [(lbl, key, _row_by_label(ws, lbl)) for lbl, key in MZAY_ROWS]
    missing = [lbl for lbl, _, r in rows if r is None]
    if missing:
        raise RuntimeError(
            "Бланк maersk_zayavka.xlsx змінено: не знайдено рядки %s%s. Заявка НЕ створена, "
            "щоб не відправити в лінію дані іншого клієнта." %
            (", ".join(missing[:3]), " та ін." if len(missing) > 3 else ""))
    for lbl, key, r in rows:
        ws.cell(row=r, column=2).value = sval(fields, key)
        nxt = MZAY_NEXT.get(key)
        if nxt and r < ws.max_row and _norm(ws.cell(row=r + 1, column=1).value) == "":
            ws.cell(row=r + 1, column=2).value = sval(fields, nxt)
    buf = io.BytesIO(); wb.save(buf); return buf.getvalue()

# ── Інформаційний лист для страхування (ТДВ «Альянс Україна») ──
INS_MAP = [
  ("BULIDING MATERIALS / БУДІВЕЛЬНІ МАТЕРІАЛИ", "cargo"),
  ("CMAU7376680", "container"),
  ("№ JMH-26-001", "!№ {invoice_no}"),
  ("11.06.2026", "invoice_date"),
  ("JE000260609був", "!{bl} був"),
  ("JE000260609", "bl"),
  ("Нагоя, Японія", "port"),
  ("на  не відбувалося", "!на {date_txt} не відбувалося"),
  ("сталися до  не будуть", "!сталися до {date_txt} не будуть"),
]
def fill_ins(fields):
    if fields.get("date") and not fields.get("date_txt"):
        fields["date_txt"] = date_dot(fields["date"])
    return fill_map("insurance.docx", INS_MAP, fields)


# ── Акт наданих послуг (Certificate of Services Rendered), двомовний бланк ──
def _uniq_cells(row):
    out, seen = [], []
    for c in row.cells:
        if any(c._tc is x for x in seen): continue
        seen.append(c._tc); out.append(c)
    return out

def _set_cell(cell, text):
    set_para(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        for r in p.runs: r.text = ""

def parse_items(raw):
    """Рядки виду 'опис | к-ть | ціна | сума' -> список позицій."""
    rows = []
    for ln in str(raw or "").split("\n"):
        ln = ln.strip()
        if not ln: continue
        parts = [x.strip() for x in ln.split("|")]
        while len(parts) < 4: parts.append("")
        rows.append(parts[:4])
    return rows

def fill_akt(fields):
    import copy
    from docx.text.paragraph import Paragraph
    d = docx.Document(os.path.join(TPL_DIR, "akt.docx"))
    if fields.get("act_date") and "-" in str(fields.get("act_date")):
        try:
            dt = datetime.date.fromisoformat(fields["act_date"])
            fields["act_date"] = dt.strftime("%B %-d, %Y")
        except Exception:
            pass
    # шапка
    hdr = _uniq_cells(d.tables[0].rows[0])[-1]
    for p in hdr.paragraphs:
        t = p.text.strip()
        if t.startswith("No.:") and fields.get("act_no"):
            set_para(p, "No.: %s" % fields["act_no"])
        elif t.startswith("Date:") and fields.get("act_date"):
            set_para(p, "Date: %s" % fields["act_date"])
    # замовник
    if fields.get("customer_block"):
        cust = _uniq_cells(d.tables[1].rows[1])[-1]
        for p in cust.paragraphs[1:]:
            p._p.getparent().remove(p._p)
        set_para(cust.paragraphs[0], fields["customer_block"])
    # рядки реквізитів
    for p in d.paragraphs:
        t = p.text.strip()
        if t.startswith("CONTRACT No:") and fields.get("contract"):
            set_para(p, "CONTRACT No: %s" % fields["contract"])
        elif t.startswith("CONTAINER") and fields.get("container"):
            set_para(p, "CONTAINER No: %s" % fields["container"])
        elif t.startswith("TO THE INVOICE No:") and fields.get("invoice_no"):
            set_para(p, "TO THE INVOICE No: %s dd %s" % (fields["invoice_no"], fields.get("invoice_date", "")))
    # позиції послуг
    items = parse_items(fields.get("items"))
    if items:
        t = d.tables[2]
        proto_tr = t.rows[1]._tr
        last_tr = t.rows[-1]._tr
        while len(t.rows) - 1 < len(items):
            last_tr.addprevious(copy.deepcopy(proto_tr))
        while len(t.rows) - 1 > len(items):
            t._tbl.remove(t.rows[-2]._tr)
        for ri, it in enumerate(items, start=1):
            cells = _uniq_cells(t.rows[ri])
            vals = [str(ri) if ri < len(items) else "%d." % ri] + it
            for ci, v in enumerate(vals):
                if ci < len(cells): _set_cell(cells[ci], v)
    if fields.get("total"):
        for row in d.tables[3].rows:
            _set_cell(_uniq_cells(row)[-1], fields["total"])
    if fields.get("amount_words"):
        for p in d.paragraphs:
            if p.text.strip().startswith("Amount in words"):
                set_para(p, "Amount in words: %s" % fields["amount_words"])
    # підпис замовника
    srow = _uniq_cells(d.tables[4].rows[0])
    if len(srow) > 1 and fields.get("customer_short"):
        set_para(srow[-1].paragraphs[0], "For and on behalf of %s:" % fields["customer_short"])
    buf = io.BytesIO(); d.save(buf); return buf.getvalue()

def clean_fn(s): return re.sub(r'[\\/:*?"<>|]', "", str(s)).strip()

TYPES = {
  "loi":         (lambda f: fill_map("loi.docx", LOI_MAP, f),      lambda f: "LOI CMA %s" % f.get("booking", "")),
  "zayavka_cma": (lambda f: fill_map("zayavka.docx", ZAYCMA_MAP, f), lambda f: "Заявка CMA %s" % f.get("booking", "")),
  "zayavka":     (fill_zp,  lambda f: "Заявка №%s %s" % (f.get("num", ""), f.get("client_short", ""))),
  "kkk":         (fill_kkk, lambda f: "Заявка ККК №%s" % f.get("num", "")),
  "t1":          (fill_t1,  lambda f: "Заявка Т1 %s" % f.get("truck", "")),
  "insurance":   (fill_ins, lambda f: "Лист страхування %s" % f.get("container", "")),
  "maersk_poa":  (fill_mpoa, lambda f: "POA Maersk %s" % f.get("bl", "")),
  "maersk_loi":  (fill_mloi, lambda f: "LOI Maersk %s" % f.get("booking", "")),
  "maersk_zayavka": (fill_mzay, lambda f: "Заявка авто Maersk %s" % f.get("booking", "")),
  "akt":         (fill_akt,  lambda f: "Akt %s %s" % (f.get("act_no", ""), f.get("customer_short", ""))),
}
# типи, чий бланк — не Word, а Excel (інше розширення файла і MIME при прикріпленні)
XLSX_TYPES = {"maersk_zayavka"}
MIME = {".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}

def get_role(jwt):
    try:
        r = urllib.request.Request(NC + "/api/v1/auth/user/me", headers={"xc-auth": jwt})
        with urllib.request.urlopen(r, timeout=10) as resp:
            email = (json.loads(resp.read().decode()).get("email") or "").lower()
        r2 = urllib.request.Request(NC + "/api/v2/tables/%s/records?limit=100" % USERS_T, headers={"xc-token": admin_tok()})
        with urllib.request.urlopen(r2, timeout=10) as resp:
            users = json.loads(resp.read().decode()).get("list", [])
        row = next((u for u in users if str(u.get("Email") or "").lower() == email and u.get("Активний") is not False), None)
        return row.get("Роль") if row else None
    except Exception:
        return None

def upload_and_attach(deal_id, fname, data, ext=".docx"):
    boundary = uuid.uuid4().hex
    body = (("--%s\r\nContent-Disposition: form-data; name=\"file\"; filename=\"%s\"\r\n"
             "Content-Type: %s\r\n\r\n") % (boundary, fname, MIME.get(ext, MIME[".docx"]))).encode() + data + ("\r\n--%s--\r\n" % boundary).encode()
    r = urllib.request.Request(NC + "/api/v2/storage/upload", data=body, method="POST",
        headers={"xc-token": admin_tok(), "Content-Type": "multipart/form-data; boundary=%s" % boundary})
    with urllib.request.urlopen(r, timeout=60) as resp:
        atts = json.loads(resp.read().decode())
    if isinstance(atts, dict): atts = [atts]
    r2 = urllib.request.Request(NC + "/api/v2/tables/%s/records?limit=1&where=%s" % (DISP_T, urllib.parse.quote("(Id,eq,%s)" % deal_id)),
        headers={"xc-token": admin_tok()})
    with urllib.request.urlopen(r2, timeout=20) as resp:
        rows = json.loads(resp.read().decode()).get("list", [])
    cur = rows[0].get("Файли") if rows else []
    if isinstance(cur, str):
        try: cur = json.loads(cur)
        except Exception: cur = []
    merged = (cur or []) + atts
    r3 = urllib.request.Request(NC + "/api/v2/tables/%s/records" % DISP_T,
        data=json.dumps([{"Id": int(deal_id), "Файли": merged}]).encode(), method="PATCH",
        headers={"xc-token": admin_tok(), "Content-Type": "application/json"})
    urllib.request.urlopen(r3, timeout=30).read()
    return merged

class H(http.server.BaseHTTPRequestHandler):
    def _send(self, code, obj):
        b = json.dumps(obj, ensure_ascii=False).encode()
        self.send_response(code); self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(b))); self.end_headers(); self.wfile.write(b)
    def do_POST(self):
        if not self.path.startswith("/gen-doc"): return self._send(404, {"error": "not found"})
        role = get_role(self.headers.get("xc-auth", ""))
        if role in DENY_ROLES: return self._send(403, {"error": "формування документів недоступне вашій ролі"})
        try:
            ln = int(self.headers.get("Content-Length", 0))
            req = json.loads(self.rfile.read(ln).decode())
            typ = req.get("type"); deal_id = req.get("dealId"); fields = req.get("fields", {})
            if typ not in TYPES or not deal_id: return self._send(400, {"error": "bad request"})
            gen, name_fn = TYPES[typ]
            data = gen(fields)
            ext = ".xlsx" if typ in XLSX_TYPES else ".docx"
            fname = clean_fn("%s %s" % (name_fn(fields), datetime.date.today().strftime("%d.%m.%Y"))) + ext
            merged = upload_and_attach(deal_id, fname, data, ext)
            self._send(200, {"ok": True, "title": fname, "files": merged})
        except Exception as e:
            self._send(500, {"error": str(e)[:150]})
    def log_message(self, *a): pass

if __name__ == "__main__":
    socketserver.TCPServer.allow_reuse_address = True
    with socketserver.TCPServer(("127.0.0.1", PORT), H) as httpd:
        httpd.serve_forever()
