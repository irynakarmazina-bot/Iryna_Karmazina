#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Збирає один документ з інструкцій репозиторію і робить .md + .docx."""
import os, re, sys

REPO = "/home/user/Iryna_Karmazina"
OUT = "/tmp/claude-0/-home-user-Iryna-Karmazina/6408fcc9-5c79-56e0-a3ea-9c1596d2bde1/scratchpad"

PARTS = [
    ("Паспорт системи", "server/PASSPORT.md"),
    ("Карта системи", "README.md"),
    ("Аварійна картка — коли щось не працює", "server/EMERGENCY.md"),
    ("Відновлення доступу", "server/ACCESS.md"),
    ("Відновлення даних з резервної копії", "server/RESTORE.md"),
]

HEAD = """# Юнітекс OS — інструкції

Складено 13.08.2026. Тримати на комп'ютері й **поза системою** — цей файл
потрібен саме тоді, коли до системи не дістатись.

Актуальна версія кожного розділу лежить у репозиторії:
https://github.com/irynakarmazina-bot/Iryna_Karmazina

| Розділ | Коли читати |
|---|---|
| 1. Паспорт системи | адреси, перші команди — почати звідси |
| 2. Карта системи | що це взагалі і як влаштоване |
| 3. Аварійна картка | **не працює**: впало, порожній екран, зламався фасад |
| 4. Відновлення доступу | **не можу зайти**: ключ, пароль, акаунт, домен |
| 5. Відновлення даних | **зникли дані**: підняти базу з копії |

---
"""


def demote(text, levels=1):
    """Опускає всі заголовки на N рівнів і прибирає перший H1 файла."""
    out, first = [], True
    for line in text.splitlines():
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            if first and len(m.group(1)) == 1:
                first = False
                continue                      # свій H1 не потрібен — є наш
            out.append("#" * min(6, len(m.group(1)) + levels) + " " + m.group(2))
        else:
            out.append(line)
    return "\n".join(out).strip()


def build_md():
    chunks = [HEAD]
    for i, (title, rel) in enumerate(PARTS, 1):
        body = open(os.path.join(REPO, rel), encoding="utf-8").read()
        chunks.append(f"\n# {i}. {title}\n\n_Джерело: `{rel}`_\n\n" + demote(body) + "\n\n---\n")
    return "\n".join(chunks)


# ── markdown → docx ────────────────────────────────────────────────────────
def to_docx(md, path):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(11)

    lines = md.splitlines()
    i, n = 0, len(lines)
    while i < n:
        ln = lines[i]

        # код
        if ln.strip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            r = p.add_run("\n".join(buf))
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            continue

        # таблиця
        if ln.strip().startswith("|") and i + 1 < n and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not re.match(r"^[\s:|-]+$", "|".join(cells)):
                    rows.append(cells)
                i += 1
            if rows:
                w = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=w)
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    cs = t.add_row().cells
                    for ci in range(w):
                        txt = clean_inline(row[ci]) if ci < len(row) else ""
                        cs[ci].text = txt
                        if ri == 0:
                            for pr in cs[ci].paragraphs:
                                for rr in pr.runs:
                                    rr.bold = True
                doc.add_paragraph()
            continue

        # заголовки
        m = re.match(r"^(#{1,6})\s+(.*)$", ln)
        if m:
            lvl = len(m.group(1))
            doc.add_heading(clean_inline(m.group(2)), level=min(lvl, 4))
            i += 1
            continue

        # горизонтальна лінія
        if ln.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.add_run("─" * 40).font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # списки
        m = re.match(r"^\s*([-*•]|\d+\.)\s+(.*)$", ln)
        if m:
            style = "List Number" if re.match(r"^\d", m.group(1)) else "List Bullet"
            add_rich(doc.add_paragraph(style=style), m.group(2))
            i += 1
            continue

        if ln.strip():
            add_rich(doc.add_paragraph(), ln.strip())
        i += 1

    doc.save(path)


def clean_inline(s):
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = re.sub(r"_(.+?)_", r"\1", s)
    s = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", s)
    return s.strip()


def add_rich(p, text):
    """Розбирає **жирний** і `код` у межах абзацу."""
    from docx.shared import Pt
    text = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", text)
    for part in re.split(r"(\*\*.+?\*\*|`.+?`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(10)
        else:
            p.add_run(part)


if __name__ == "__main__":
    md = build_md()
    md_path = os.path.join(OUT, "Юнітекс-OS-інструкції.md")
    open(md_path, "w", encoding="utf-8").write(md)
    docx_path = os.path.join(OUT, "Юнітекс-OS-інструкції.docx")
    to_docx(md, docx_path)
    print("md:", md_path, len(md.splitlines()), "рядків")
    print("docx:", docx_path, os.path.getsize(docx_path), "байт")
